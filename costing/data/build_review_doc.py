#!/usr/bin/env python3
"""Build the short, forward-ready Word review doc for the team."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
# base style
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

NAVY = RGBColor(0x30,0x54,0x96)
RED  = RGBColor(0xC0,0x00,0x00)
GREEN= RGBColor(0x2E,0x7D,0x32)
GREY = RGBColor(0x60,0x60,0x60)

def h1(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p
def h2(t, color=NAVY):
    p = doc.add_paragraph(); p.space_before = Pt(8); r = p.add_run(t); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = color
    return p
def body(t, italic=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic
    if color: r.font.color.rgb = color
    return p
def bullet(t, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(t)
    else:
        p.add_run(t)
    return p

# Title block
h1("Belt Costing Software — Design Review")
body("Review of flow.docx (system working design) against the agreed specification.", italic=True, color=GREY)
body("Date: 28 May 2026   ·   Prepared for the development team   ·   Status of flow.docx: work-in-progress", italic=True, color=GREY)
doc.add_paragraph()

# Verdict
h2("Verdict — strong start, ~80% aligned")
body("The working design faithfully captures our current dealer costing sheet — field layout, override pattern, "
     "CD/VD pricing, revisions and PO tracking are all correct. Before building further, two structural changes are "
     "needed so the system delivers the flexibility we set out to achieve and can later connect to the Ravasco "
     "Formulations app. The single most important one: the selectable skim compound (our original requirement) is not yet in the design.")

# Matches
h2("What already matches — keep as designed", GREEN)
for t in [
  "Belt Rating built by concatenation (fabric + strength + ply).",
  "Product Type → Conveyor Type from masters.",
  "Customer-driven header; special note from the enquiry email.",
  "Open-End / Endless construction toggle.",
  "Override-everything pattern (blank = use master, else use the typed value).",
  "CD and VD pricing shown side by side.",
  "Per-meter weights and costs.",
  "Revisions 1–4, Order Rate, Actual GP, PO Status.",
  "Multiple fabric suppliers (MIT / SRF / others).",
]:
    bullet(t)

# Adopt
h2("Good ideas in the design we will ADOPT into the spec", GREEN)
for lead, t in [
  ("Metric/Imperial unit toggle — ", "we under-specified this; add it."),
  ("Calculated vs Actual belt weight — ", "keep the manual actual-weight entry."),
  ("Floating / Non-Floating breaker — ", "add as a breaker attribute (good catch)."),
  ("Floor Price line — ", "surface the break-even price clearly for sales."),
]:
    bullet(t, bold_lead=lead)

# Changes needed
h2("Important changes before building further", RED)
p = doc.add_paragraph(); r = p.add_run("1.  Make SKIM a selectable compound (the original requirement)."); r.bold = True; r.font.color.rgb = RED
body("Today the carcass skim rate is bundled inside the Grade dropdown and only appears as a manual override. "
     "It must be a first-class compound you select from the compound list (role = SKIM), with recommended (matching) "
     "skims listed first and a warning if a non-matching skim is chosen. This is the exact feature we set out to add "
     "(“option of selecting a different skim compound if required”).")

p = doc.add_paragraph(); r = p.add_run("2.  Merge Grade_master + skim master into ONE Compound Master with a “role” field."); r.bold = True; r.font.color.rgb = RED
body("Each compound carries one or more roles (Top Cover, Bottom Cover, Skim, Sidewall, Cleat, …). This single change "
     "unlocks: selectable skim, separable top/bottom cover, controlled cross-use, and — most importantly — a clean path "
     "to pull live prices from the Ravasco formulation app (where a compound’s ₹/kg is computed from its recipe). "
     "Keeping rates bundled in Grade means prices stay typed-in forever.")

p = doc.add_paragraph(); r = p.add_run("3.  Allow different Top vs Bottom cover compounds."); r.bold = True; r.font.color.rgb = RED
body("Default both to the same; allow either to be changed (a cheaper bottom cover is common).")

p = doc.add_paragraph(); r = p.add_run("4.  Add Edge Type (Cut-Edge / Moulded)."); r.bold = True; r.font.color.rgb = RED
body("It drives the width-wastage calculation (Cut-Edge +30 mm, Moulded +0). Not present in the current design.")

# Resolved
h2("Confirmed (no change needed)")
for lead, t in [
  ("Multi-belt quotes — ", "one quotation holds many belt lines; “Sr. No” is the line number. (Agreed.)"),
  ("Rate freeze on send — ", "to be added: snapshot all rates when a quote is sent, so master changes never rewrite old quotes."),
  ("Width from Width_Master — ", "kept as a dropdown to avoid manual-entry errors (custom widths added as master rows)."),
  ("Per-belt-type fields — ", "show only the fields applicable to the selected belt type. (Agreed.)"),
]:
    bullet(t, bold_lead=lead)

# Handover
h2("Supporting material provided")
for t in [
  "Full specification — costing-app-planning/00-START-HERE.md (read first; glossary + all decisions).",
  "Compound master + skim↔cover matching — docs 02, 02b, 02c.",
  "Database schema (Prisma, Ravasco-aligned) — schema/schema.prisma (validated).",
  "Master import templates, pre-seeded — import-templates/Belt-Costing-Master-Import-Templates.xlsx.",
  "Corrected, master-driven sample sheet — Corrected-Multiply-Belt-Costing.xlsx (reproduces the original to the rupee).",
]:
    bullet(t)

doc.add_paragraph()
body("In short: keep the excellent field-level UI; upgrade the master layer underneath (one Compound Master with roles, "
     "selectable skim, edge type). That is what turns this from a digitised sheet into the flexible, single-source-of-truth "
     "system we want.", italic=True, color=GREY)

out = os.path.join(os.path.dirname(__file__), "Belt-Costing-Design-Review.docx")
doc.save(out)
print("Saved:", out)
